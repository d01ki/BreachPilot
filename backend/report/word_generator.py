"""
Word Report Generator - Professional Word document generation
"""

from pathlib import Path
from datetime import datetime
from typing import Dict, Any
import logging
import tempfile
import subprocess
import json

logger = logging.getLogger(__name__)

class WordReportGenerator:
    """Generate professional Word documents with advanced formatting"""
    
    def __init__(self, reports_dir: Path):
        self.reports_dir = reports_dir
        self.reports_dir.mkdir(exist_ok=True)
    
    def generate_word_report(self, target_ip: str, report_data: Dict[str, Any], output_path: Path) -> Path:
        """Generate comprehensive Word report"""
        logger.info(f"Generating Word report for {target_ip}")
        
        try:
            # Try to use python-docx if available
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
            from docx.oxml.shared import OxmlElement, qn
            
            self._generate_docx_report(target_ip, report_data, output_path)
            
        except ImportError:
            logger.warning("python-docx not available, using fallback method")
            self._generate_fallback_word(target_ip, report_data, output_path)
            
        except Exception as e:
            logger.error(f"Word generation failed: {e}")
            self._generate_error_word(output_path, target_ip, str(e))
        
        return output_path
    
    def _generate_docx_report(self, target_ip: str, report_data: Dict[str, Any], output_path: Path):
        """Generate DOCX report using python-docx"""
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        
        doc = Document()
        
        # Set up document styles
        self._setup_document_styles(doc)
        
        # Cover page
        self._add_cover_page(doc, target_ip, report_data)
        
        # Table of contents
        self._add_table_of_contents(doc)
        
        # Executive summary
        self._add_executive_summary(doc, report_data)
        
        # Risk assessment
        self._add_risk_assessment(doc, report_data)
        
        # Technical analysis
        self._add_technical_analysis(doc, report_data)
        
        # Recommendations
        self._add_recommendations(doc, report_data)
        
        # Remediation roadmap
        self._add_remediation_roadmap(doc, report_data)
        
        # Compliance analysis
        self._add_compliance_analysis(doc, report_data)
        
        # Footer
        self._add_footer(doc, report_data)
        
        # Save document
        doc.save(str(output_path))
        logger.info(f"Word report generated successfully: {output_path}")
    
    def _setup_document_styles(self, doc):
        """Set up document styles"""
        from docx.shared import Pt
        
        # Title style
        title_style = doc.styles.add_style('CustomTitle', 1)
        title_font = title_style.font
        title_font.name = 'Arial'
        title_font.size = Pt(24)
        title_font.bold = True
        
        # Heading 1 style
        h1_style = doc.styles['Heading 1']
        h1_font = h1_style.font
        h1_font.name = 'Arial'
        h1_font.size = Pt(18)
        h1_font.bold = True
        
        # Heading 2 style
        h2_style = doc.styles['Heading 2']
        h2_font = h2_style.font
        h2_font.name = 'Arial'
        h2_font.size = Pt(14)
        h2_font.bold = True
        
        # Normal style
        normal_style = doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Arial'
        normal_font.size = Pt(11)
    
    def _add_cover_page(self, doc, target_ip: str, report_data: Dict[str, Any]):
        """Add cover page"""
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Add title
        title = doc.add_heading('SECURITY ASSESSMENT REPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add subtitle
        subtitle = doc.add_paragraph(f'Target: {target_ip}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.size = Pt(16)
        subtitle_run.font.bold = True
        
        # Add metadata
        metadata = report_data.get("report_metadata", {})
        
        metadata_para = doc.add_paragraph()
        metadata_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        metadata_para.add_run(f"Report ID: {metadata.get('report_id', 'N/A')}\n")
        metadata_para.add_run(f"Date: {metadata.get('generation_date', 'N/A')}\n")
        metadata_para.add_run(f"Classification: {metadata.get('classification', 'CONFIDENTIAL')}\n")
        
        # Add page break
        doc.add_page_break()
    
    def _add_table_of_contents(self, doc, report_data: Dict[str, Any]):
        """Add table of contents"""
        toc_heading = doc.add_heading('Table of Contents', level=1)
        
        toc_items = [
            "Executive Summary",
            "Risk Assessment", 
            "Technical Analysis",
            "Detailed Recommendations",
            "Remediation Roadmap",
            "Compliance Analysis",
            "Appendices"
        ]
        
        for item in toc_items:
            para = doc.add_paragraph()
            para.add_run(f"• {item}")
        
        doc.add_page_break()
    
    def _add_executive_summary(self, doc, report_data: Dict[str, Any]):
        """Add executive summary section"""
        doc.add_heading('Executive Summary', level=1)
        
        executive_summary = report_data.get("executive_summary", {})
        
        # Assessment overview
        doc.add_heading('Assessment Overview', level=2)
        overview_para = doc.add_paragraph(executive_summary.get('assessment_overview', 'Assessment completed.'))
        
        # Key findings
        doc.add_heading('Key Findings', level=2)
        key_findings = executive_summary.get('key_findings', [])
        for finding in key_findings:
            para = doc.add_paragraph(finding, style='List Bullet')
        
        # Critical risks
        doc.add_heading('Critical Risks', level=2)
        critical_risks = executive_summary.get('critical_risks', [])
        for risk in critical_risks:
            para = doc.add_paragraph(risk, style='List Bullet')
        
        # Immediate actions
        doc.add_heading('Immediate Actions Required', level=2)
        immediate_actions = executive_summary.get('immediate_actions', [])
        for action in immediate_actions:
            para = doc.add_paragraph(action, style='List Bullet')
    
    def _add_risk_assessment(self, doc, report_data: Dict[str, Any]):
        """Add risk assessment section"""
        doc.add_heading('Risk Assessment', level=1)
        
        risk_assessment = report_data.get("risk_assessment", {})
        
        # Risk metrics table
        doc.add_heading('Risk Metrics', level=2)
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Value'
        
        # Add risk data
        risk_data = [
            ('Overall Risk Level', risk_assessment.get('overall_risk_level', 'Unknown')),
            ('Risk Score', f"{risk_assessment.get('risk_score', 0):.1f}/10"),
            ('Critical Vulnerabilities', str(risk_assessment.get('critical_vulnerabilities', 0))),
            ('High Vulnerabilities', str(risk_assessment.get('high_vulnerabilities', 0))),
            ('Financial Impact', risk_assessment.get('financial_impact_estimate', 'Unknown')),
            ('Remediation Priority', risk_assessment.get('remediation_priority', 'Unknown'))
        ]
        
        for metric, value in risk_data:
            row_cells = table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = value
        
        # Business impact
        doc.add_heading('Business Impact Analysis', level=2)
        impact_para = doc.add_paragraph(f"Impact: {risk_assessment.get('business_impact', 'Assessment in progress.')}")
        technical_para = doc.add_paragraph(f"Technical Impact: {risk_assessment.get('technical_impact', 'Assessment in progress.')}")
        likelihood_para = doc.add_paragraph(f"Likelihood: {risk_assessment.get('likelihood', 'Assessment in progress.')}")
    
    def _add_technical_analysis(self, doc, report_data: Dict[str, Any]):
        """Add technical analysis section"""
        doc.add_heading('Technical Analysis', level=1)
        
        technical_analysis = report_data.get("technical_analysis", {})
        
        # Attack vectors
        doc.add_heading('Attack Vectors', level=2)
        
        attack_vectors = technical_analysis.get('attack_vectors', [])
        if attack_vectors:
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            
            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'CVE ID'
            hdr_cells[1].text = 'Attack Type'
            hdr_cells[2].text = 'Target Service'
            hdr_cells[3].text = 'Complexity'
            hdr_cells[4].text = 'Impact'
            
            # Add attack vector data
            for vector in attack_vectors[:10]:  # Limit to top 10
                row_cells = table.add_row().cells
                row_cells[0].text = vector.get('cve_id', 'N/A')
                row_cells[1].text = vector.get('attack_type', 'Unknown')
                row_cells[2].text = vector.get('target_service', 'Unknown')
                row_cells[3].text = vector.get('complexity', 'Unknown')
                row_cells[4].text = vector.get('impact', 'Unknown')
        
        # Lateral movement analysis
        doc.add_heading('Lateral Movement Analysis', level=2)
        lateral_movement = technical_analysis.get('lateral_movement_potential', {})
        
        doc.add_paragraph(f"Network Segments: {lateral_movement.get('network_segments', 'Unknown')}")
        doc.add_paragraph(f"Privileged Services: {lateral_movement.get('privileged_services', 'Unknown')}")
        doc.add_paragraph(f"Risk Level: {lateral_movement.get('risk_level', 'Unknown')}")
    
    def _add_recommendations(self, doc, report_data: Dict[str, Any]):
        """Add recommendations section"""
        doc.add_heading('Detailed Recommendations', level=1)
        
        recommendations = report_data.get("detailed_recommendations", [])
        
        for i, rec in enumerate(recommendations, 1):
            doc.add_heading(f"{i}. {rec.get('title', 'Recommendation')}", level=2)
            
            # Recommendation details
            details = [
                ('Priority', rec.get('priority', 'Unknown')),
                ('Category', rec.get('category', 'Unknown')),
                ('Description', rec.get('description', 'No description available.')),
                ('Business Justification', rec.get('business_justification', 'No justification provided.')),
                ('Timeline', rec.get('timeline', 'Unknown')),
                ('Estimated Cost', rec.get('estimated_cost', 'Unknown')),
                ('Success Metrics', rec.get('success_metrics', 'Unknown'))
            ]
            
            for label, value in details:
                para = doc.add_paragraph()
                para.add_run(f"{label}: ").bold = True
                para.add_run(str(value))
    
    def _add_remediation_roadmap(self, doc, report_data: Dict[str, Any]):
        """Add remediation roadmap section"""
        doc.add_heading('Remediation Roadmap', level=1)
        
        roadmap = report_data.get("remediation_roadmap", {})
        
        phases = roadmap.get("phases", [])
        for phase in phases:
            doc.add_heading(phase.get('phase', 'Phase'), level=2)
            
            # Phase details
            doc.add_paragraph(f"Focus: {phase.get('focus', 'Unknown')}")
            
            doc.add_paragraph("Activities:")
            activities = phase.get('activities', [])
            for activity in activities:
                doc.add_paragraph(activity, style='List Bullet')
            
            doc.add_paragraph(f"Success Criteria: {phase.get('success_criteria', 'Unknown')}")
            doc.add_paragraph(f"Budget: {phase.get('budget', 'Unknown')}")
        
        # Summary
        doc.add_heading('Roadmap Summary', level=2)
        summary = [
            ('Total Estimated Cost', roadmap.get('total_estimated_cost', 'Unknown')),
            ('Expected ROI', roadmap.get('expected_roi', 'Unknown')),
            ('Timeline', roadmap.get('timeline', 'Unknown'))
        ]
        
        for label, value in summary:
            para = doc.add_paragraph()
            para.add_run(f"{label}: ").bold = True
            para.add_run(str(value))
    
    def _add_compliance_analysis(self, doc, report_data: Dict[str, Any]):
        """Add compliance analysis section"""
        doc.add_heading('Compliance Analysis', level=1)
        
        compliance_analysis = report_data.get("compliance_analysis", {})
        
        # Applicable frameworks
        doc.add_heading('Applicable Frameworks', level=2)
        frameworks = compliance_analysis.get('applicable_frameworks', [])
        for framework in frameworks:
            doc.add_paragraph(framework, style='List Bullet')
        
        # Compliance gaps
        doc.add_heading('Compliance Gaps', level=2)
        gaps = compliance_analysis.get('compliance_gaps', [])
        for gap in gaps:
            doc.add_heading(gap.get('framework', 'Unknown Framework'), level=3)
            doc.add_paragraph(f"Requirement: {gap.get('requirement', 'Unknown')}")
            doc.add_paragraph(f"Gap: {gap.get('gap', 'Unknown')}")
            doc.add_paragraph(f"Impact: {gap.get('impact', 'Unknown')}")
    
    def _add_footer(self, doc, report_data: Dict[str, Any]):
        """Add footer information"""
        doc.add_heading('Document Information', level=1)
        
        metadata = report_data.get("report_metadata", {})
        
        footer_info = [
            ('Generated by', 'BreachPilot Professional Security Assessment Platform'),
            ('Report Classification', metadata.get('classification', 'CONFIDENTIAL')),
            ('Generated', datetime.now().strftime('%Y-%m-%d %H:%M:%S UTC')),
            ('Validity Period', metadata.get('validity_period', '90 days'))
        ]
        
        for label, value in footer_info:
            para = doc.add_paragraph()
            para.add_run(f"{label}: ").bold = True
            para.add_run(str(value))
    
    def _generate_fallback_word(self, target_ip: str, report_data: Dict[str, Any], output_path: Path):
        """Generate fallback Word document"""
        logger.info("Generating fallback Word document")
        
        # Create HTML content first
        html_content = self._generate_html_content(target_ip, report_data)
        
        try:
            # Try to use pandoc to convert HTML to DOCX
            with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False) as temp_file:
                temp_file.write(html_content)
                temp_file.flush()
                
                # Convert to DOCX using pandoc
                subprocess.run([
                    'pandoc', temp_file.name,
                    '-o', str(output_path),
                    '--reference-doc=/usr/share/pandoc/data/reference.docx'  # Use system reference if available
                ], check=True, capture_output=True)
                
                import os
                os.unlink(temp_file.name)
                logger.info(f"Fallback Word document generated successfully: {output_path}")
                
        except (subprocess.CalledProcessError, FileNotFoundError):
            # Final fallback: create a simple HTML file with .docx extension
            logger.warning("Pandoc not available, creating HTML-based document")
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
    
    def _generate_html_content(self, target_ip: str, report_data: Dict[str, Any]) -> str:
        """Generate HTML content for Word conversion"""
        # This is a simplified version of the HTML content
        # The full implementation would be similar to the PDF generator
        return f"""
        <html>
        <head>
            <title>Security Assessment Report - {target_ip}</title>
        </head>
        <body>
            <h1>Security Assessment Report - {target_ip}</h1>
            <p>Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
            <h2>Executive Summary</h2>
            <p>{report_data.get('executive_summary', {}).get('assessment_overview', 'Assessment completed.')}</p>
            <h2>Risk Assessment</h2>
            <p>Risk Level: {report_data.get('risk_assessment', {}).get('overall_risk_level', 'Unknown')}</p>
            <p>Risk Score: {report_data.get('risk_assessment', {}).get('risk_score', 0):.1f}/10</p>
        </body>
        </html>
        """
    
    def _generate_error_word(self, output_path: Path, target_ip: str, error_msg: str):
        """Generate error Word document"""
        error_content = f"""
        <html>
        <head><title>Security Assessment Report - Error</title></head>
        <body>
            <h1>Security Assessment Report - ERROR</h1>
            <h2>Target: {target_ip}</h2>
            <p>Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
            <h2>Word Generation Failed</h2>
            <p>Error: {error_msg}</p>
            <p>Please contact the system administrator or check the HTML report instead.</p>
            <p>Generated by BreachPilot Professional Security Assessment Platform</p>
        </body>
        </html>
        """
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(error_content)
